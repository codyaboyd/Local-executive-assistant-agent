from pathlib import Path

import pytest

from app.tools import docx


def test_build_docx_chunks_includes_required_metadata(monkeypatch) -> None:
    def fake_extract(path: Path):
        return [
            docx.DOCXSection("Travel Policy", "Travel Policy\n\nReceipts are required."),
            docx.DOCXSection("Approvals", "Approvals are due Friday."),
        ]

    monkeypatch.setattr(docx, "extract_docx_sections", fake_extract)

    chunks = docx.build_docx_chunks("/tmp/handbook.docx")

    assert [chunk.text for chunk in chunks] == ["Travel Policy\n\nReceipts are required.", "Approvals are due Friday."]
    assert chunks[0].metadata == {
        "source": "handbook.docx",
        "section_heading": "Travel Policy",
        "file_type": "docx",
        "section": 1,
        "chunk": 1,
    }
    assert chunks[1].metadata == {
        "source": "handbook.docx",
        "section_heading": "Approvals",
        "file_type": "docx",
        "section": 2,
        "chunk": 1,
    }


def test_ingest_docx_stores_chunks_with_metadata(monkeypatch) -> None:
    monkeypatch.setattr(
        docx,
        "build_docx_chunks",
        lambda path: [
            docx.DOCXChunk(
                "Travel policy requires receipts.",
                {"source": "handbook.docx", "section_heading": "Travel Policy", "file_type": "docx"},
            )
        ],
    )

    class FakeVectorStore:
        def __init__(self) -> None:
            self.added = None

        def add_documents(self, chunks, metadata):
            self.added = (chunks, metadata)

    store = FakeVectorStore()

    assert docx.ingest_docx("handbook.docx", store) == 1
    assert store.added == (
        ["Travel policy requires receipts."],
        [{"source": "handbook.docx", "section_heading": "Travel Policy", "file_type": "docx"}],
    )


def test_extract_docx_sections_rejects_non_docx(tmp_path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("not a docx")

    with pytest.raises(ValueError):
        docx.extract_docx_sections(path)


def test_extract_docx_sections_reads_headings_paragraphs_and_tables(tmp_path) -> None:
    from docx import Document

    path = tmp_path / "handbook.docx"
    document = Document()
    document.add_heading("Travel Policy", level=1)
    document.add_paragraph("Receipts are required.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Limit"
    table.cell(0, 1).text = "$50"
    table.cell(1, 0).text = "Approval"
    table.cell(1, 1).text = "Manager"
    document.add_heading("Approvals", level=2)
    document.add_paragraph("Submit by Friday.")
    document.save(path)

    sections = docx.extract_docx_sections(path)

    assert sections[0].heading == "Travel Policy"
    assert "Receipts are required." in sections[0].text
    assert "Limit | $50" in sections[0].text
    assert sections[1] == docx.DOCXSection("Approvals", "Approvals\n\nSubmit by Friday.")
